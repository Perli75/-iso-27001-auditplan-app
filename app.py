# ─────────────────────────────────────────────────────────────────────────────
# ISO 27001 Audit-Plan & Work-Pack Generator • app.py  (v3.0 • 07-May-2025)
# Produces TWO Word files per run:
#   •  <Client>-Audit-Plan-[Stage|Surveillance].docx
#   •  <Client>-Work-Pack.docx
#
# Requires:  streamlit  |  python-docx  |  (logo optional)
# ─────────────────────────────────────────────────────────────────────────────

import math
from datetime import date, timedelta
from io import BytesIO
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from utils import calculate_audit_days

# ── helpers ──────────────────────────────────────────────────────────────────
def std_font(doc, size=11):
    font = doc.styles["Normal"].font
    font.name, font.size = "Calibri", Pt(size)

def add_h1(doc, txt):
    h = doc.add_heading(txt, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h2(doc, txt):
    doc.add_heading(txt, level=2)

def add_key_table(doc, data):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in data.items():
        r = tbl.add_row().cells
        r[0].text, r[1].text = k, v
        r[0].paragraphs[0].runs[0].font.bold = True
        r[0].vertical_alignment = r[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph()

def bullets(doc, lines):
    for l in lines:
        doc.add_paragraph(l, style="List Bullet")

def timetable(doc, start_dt, days_span, lead, contact):
    add_h2(doc, "Audit Schedule (Draft)")
    slots = [("09:00-12:00", "Opening meeting, context & leadership review"),
             ("13:00-16:30", "Controls verification, evidence collection")]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style, tbl.alignment = "Table Grid", WD_TABLE_ALIGNMENT.CENTER
    hdr = ["Date", "Time", "Auditor", "Activity", "Attendees"]
    for i, h in enumerate(hdr):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
    for d in range(days_span):
        cur = (start_dt + timedelta(days=d)).strftime("%d %b %Y")
        for tm, act in slots:
            r = tbl.add_row().cells
            r[0].text, r[1].text = cur, tm
            r[2].text = lead
            r[3].text = act
            r[4].text = contact
    doc.add_paragraph()

def save_to_button(filename, document):
    buf = BytesIO()
    document.save(buf); buf.seek(0)
    st.download_button(f"⬇️ Download {filename}", buf,
                       file_name=filename,
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ── page config ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="ISO 27001 Audit Generator", page_icon="🗂️")
st.title("ISO 27001 Audit-Plan & Work-Pack Generator")
try:
    st.image("Logo.png", width=130)
except FileNotFoundError:
    pass
st.caption("Fill in the details → click **Create** → download two Word files ready to send.")

# ── form inputs ──────────────────────────────────────────────────────────────
with st.form("planform", clear_on_submit=False):
    colA, colB = st.columns(2)
    with colA:
        client   = st.text_input("Client name", "Skyline")
        scope    = st.text_area("Scope of certification", "Provision of XYZ managed services", height=80)
        contact  = st.text_input("Client primary contact", "Jane Doe – CISO")
        address  = st.text_input("Primary site address", "123 Cyber St, Brisbane QLD 4000")
    with colB:
        audit_type = st.selectbox("Audit type", ["Stage 1", "Stage 2", "Surveillance", "Recertification"], index=2)
        start_dt   = st.date_input("Audit start date", value=date.today())
        days_span  = st.number_input("Audit duration (calendar days)", 1, 10, value=2)
        lead       = st.text_input("Lead auditor", "Perla Chandler")
        team       = st.text_input("Additional auditors / SMEs", "Michelle Coleman")

    st.markdown("---")
    st.subheader("Objectives (one per line)")
    objectives = st.text_area("", "\n".join([
        "Confirm that the ISMS conforms with ISO/IEC 27001:2022.",
        "Verify statutory, regulatory and contractual compliance.",
        "Evaluate the effectiveness of implemented controls.",
        "Identify opportunities for improvement."
    ]), height=120)

    st.subheader("Documents / Records requested (one per line)")
    docs_req = st.text_area("", "Management Review minutes\nInternal audit reports\nRisk Register\nStatement of Applicability", height=100)

    employees = st.number_input("Employees in scope", 1, 50_000, 100)
    sites     = st.number_input("Physical locations", 1, 50, 1)

    submitted = st.form_submit_button("Create Word Audit-Plan & Work-Pack")

# ── generate on submit ───────────────────────────────────────────────────────
if submitted:
    # 1. Effort calculation (simple rule – adjust in utils.py)
    base_d, extra_h = calculate_audit_days(employees, sites)
    FACTOR = {"Stage 1":0.4, "Stage 2":1.0, "Surveillance":0.3, "Recertification":0.6}[audit_type]
    total_days = round(base_d*FACTOR + (extra_h/8)*FACTOR, 2)
    end_dt = start_dt + timedelta(days=days_span-1)

    # 2. Build Audit-Plan -----------------------------------------------------
    plan = Document(); std_font(plan)
    add_h1(plan, f"{client} – ISO 27001 {audit_type} Audit Plan")
    add_key_table(plan, {
        "Audit criteria": "ISO/IEC 27001:2022",
        "Scope": scope,
        "Site": address,
        "Audit dates": f"{start_dt.strftime('%d %b %Y')} – {end_dt.strftime('%d %b %Y')}",
        "Audit team": f"{lead} (Lead)\n{team}",
        "Client contact": contact,
        "Calculated effort": f"{total_days} audit-days",
    })
    add_h2(plan, "Audit Objectives")
    bullets(plan, [o.strip() for o in objectives.splitlines() if o.strip()])

    add_h2(plan, "Documentation / Records Requested")
    bullets(plan, [d.strip() for d in docs_req.splitlines() if d.strip()])

    plan.add_page_break()
    timetable(plan, start_dt, days_span, lead, contact)

    plan.add_page_break()
    add_h2(plan, "Closing Meeting & Reporting")
    plan.add_paragraph(
        "The closing meeting will be held on the final day at 16:30 local time. "
        "Audit findings, non-conformances and timelines for corrective action "
        "will be discussed. The audit report will be issued within seven working days."
    )

    # 3. Build Work-Pack ------------------------------------------------------
    wp = Document(); std_font(wp)
    add_h1(wp, f"Audit Briefing & Work-Pack – {client}")
    wp.add_paragraph(
        "Purpose: to brief the audit team on agreed location, scope, criteria, "
        "timeframes and required evidence before commencing the audit."
    )

    add_h2(wp, "Client & Audit Details")
    add_key_table(wp, {
        "Client": client,
        "Audit type": audit_type,
        "Criteria": "ISO/IEC 27001:2022",
        "Scope": scope,
        "Dates": f"{start_dt.strftime('%d %b %Y')} – {end_dt.strftime('%d %b %Y')}",
        "Team": f"{lead} (Lead)\n{team}",
    })

    add_h2(wp, "High-Level Schedule")
    timetable(wp, start_dt, days_span, lead, contact)

    add_h2(wp, "Evidence Checklist")
    bullets(wp, [d.strip() for d in docs_req.splitlines() if d.strip()])

    wp.add_paragraph(); wp.add_paragraph("Prepared by: …………………………………")  # signature line

    # 4. Stream files ---------------------------------------------------------
    st.success("Files generated!  Download below:")
    save_to_button(f"{client.replace(' ','_')}-Audit-Plan-{audit_type}.docx", plan)
    save_to_button(f"{client.replace(' ','_')}-Work-Pack.docx", wp)
